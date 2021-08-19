from selenium import webdriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By
import time
import os
from bs4 import BeautifulSoup
from docx import Document
import datetime
import re
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


class RiskLevel:
    HIGH = 1
    MEDIUM = 2
    LOW = 3

    HIGH_TAG = u"高风险"
    MEDIUM_TAG = u"中风险"
    LOW_TAG = u"低风险"

    HIGH_LV_TEXT = u"高风险区"
    MEDIUM_LV_TEXT = u"中风险区"
    LOW_LV_TEXT = u"低风险区"

    text_2_level = {
        HIGH_TAG: HIGH,
        MEDIUM_TAG: MEDIUM,
        LOW_TAG: LOW,
    }

    level_2_text = {
        HIGH: HIGH_LV_TEXT,
        MEDIUM: MEDIUM_LV_TEXT,
        LOW: LOW_LV_TEXT,
    }

    @staticmethod
    def text_to_risk_level(risk_str):
        for tag, level in RiskLevel.text_2_level.items():
            if risk_str.find(tag) >= 0:
                return level
        raise Exception("unkown risk str:" + risk_str)

    @staticmethod
    def risk_level_to_text(risk_lv):
        return RiskLevel.level_2_text.get(risk_lv)


class RiskBlock:

    def __init__(self, province, block):
        self.province = province
        self.block = block
        self.is_new_add = False

    @property
    def full_name(self):
        return self.province + self.block

    def full_name_with_sep(self):
        return self.province + " " + self.block

    def full_name_with_new_tag(self):
        new_tag = SpiderBase.NEW_ADD_TAG if self.is_new_add else ""
        return "{0}{1}{2}".format(self.province, self.block, new_tag)

    def block_name_with_new_tag(self):
        new_tag = SpiderBase.NEW_ADD_TAG if self.is_new_add else ""
        return "{0}{1}".format(self.block, new_tag)

    def __eq__(self, other):
        return self.full_name == other.full_name
        # return self.province == other.province and self.block == other.block


class SpiderBase(object):
    NEW_ADD_TAG = u"(新增)"
    OUTPUT_PATH = "../"
    TEMP_PATH = "../tmp/"
    TEMPLATE_PATH = "../template/"

    def __init__(self, show_browser):
        self.risk_blocks = {RiskLevel.HIGH: [], RiskLevel.MEDIUM: []}
        self.last_blocks = None
        self.result_time = None
        self._wait_time = 1.5
        self._init_driver(show_browser)

    def _init_driver(self, show_browser):
        if show_browser:
            # 创建可见的Chrome浏览器， 方便调试
            self.driver = webdriver.Chrome()
        else:
            # 创建Chrome的无头浏览器
            opt = webdriver.ChromeOptions()
            opt.headless = True
            self.driver = webdriver.Chrome(options=opt)

        self.driver.implicitly_wait(10)

    def sleep(self):
        if self._wait_time <= 0:
            return
        time.sleep(self._wait_time)

    def output_html(self):
        soup = BeautifulSoup(self.driver.page_source, 'html.parser')
        print(soup.prettify())

    def set_result_time(self, time_txt):
        if self.result_time is not None:
            return
        match_obj = re.match(r'\D*(\d+-\d+-\d+ \d+).*', time_txt)
        time_str = match_obj.group(1) + ":00:00"
        self.result_time = datetime.datetime.strptime(time_str, "%Y-%m-%d %H:%M:%S")

    def add_risk_block(self, risk_text, city_name, block_name):
        # print("add_risk_block", risk_text, city_name, block_name)
        risk_level = RiskLevel.text_to_risk_level(risk_text)
        if risk_level == RiskLevel.LOW:
            return
        # block_lst = block_name.split(" ")
        risk_block = RiskBlock(city_name, block_name)
        # block_name = block_name.replace(" ", "")
        lst_block = self.risk_blocks.get(risk_level)
        if risk_block in lst_block:
            print(u"小区{0}已存在{1}中".format(block_name, RiskLevel.risk_level_to_text(risk_level)))
            return
        lst_block.append(risk_block)

    def format_result_time(self):
        time_tup = self.result_time.timetuple()
        return u"{0}月{1}日{2}时".format(time_tup.tm_mon, time_tup.tm_mday, time_tup.tm_hour)

    def format_result_time_ex(self):
        return self.result_time.strftime("%Y-%m-%d-%H")

    def format_last_result_time(self):
        time_tup = datetime.datetime.fromtimestamp(self.result_time.timestamp() - 24 * 3600).timetuple()
        return u"{0}月{1}日{2}时".format(time_tup.tm_mon, time_tup.tm_mday, time_tup.tm_hour)

    # @staticmethod
    def get_last_result_file_name(self):
        now_time = time.time()
        # time_info = time.localtime(now_time)
        cur_time_tup = self.result_time.timetuple()
        last_time_tup = time.localtime(now_time - 3600 * 24)
        # last_str_time = time.strftime("%Y-%m-%d", time.localtime(now_time - 3600 * 24))
        # last_f_name = u"{0}RiskBlock_{1}.txt".format(SpiderBase.TEMP_PATH, last_str_time)
        # time_str = self.format_last_result_time()

        files = []
        def _parse_file_name(path, f_name):
            # print(f_name)
            if path != SpiderBase.OUTPUT_PATH:
                return
            if not f_name.endswith(".doc") and not f_name.endswith(".docx"):
                return
            pattern = re.compile(r'.*(\d+)月(\d+)日(\d+)时.*')
            m = pattern.match(f_name)
            if m is None:
                print("{0} 文件名里不包含截止时间,某月某日某时!!!!".format(f_name))
                return
            month = int(m[1])
            day = int(m[2])
            hour = int(m[3])
            if (month != cur_time_tup.tm_mon or day != cur_time_tup.tm_mday) and \
                    (month != last_time_tup.tm_mon or day != last_time_tup.tm_mday):
                return
            if month == cur_time_tup.tm_mon and day == cur_time_tup.tm_mday and hour == cur_time_tup.tm_hour:
                return
            # print(m[0], type(m[1]), m[2], m[3])
            files.append(f_name)

        SpiderBase.walk_dir(SpiderBase.OUTPUT_PATH, _parse_file_name)
        print(files)
        if not files:
            return
        files.sort()
        return SpiderBase.OUTPUT_PATH + files[len(files)-1]

    def init_last_risk_block_lst(self):
        last_f_name = self.get_last_result_file_name()
        if last_f_name is None or not os.path.exists(last_f_name):
            print("!!!无法输出新增记录,前一天的记录不存在!!!{0}\n".format(last_f_name))
            return
        print(u"上一次通知文件:{0}".format(last_f_name))
        # self.last_blocks = self.read_record_from_file(last_f_name)
        self.last_blocks = self.read_record_from_word(last_f_name)
        print("昨日({0})高风险地区:{1}个, 中风险地区:{2}个".format(self.format_last_result_time(),
                                                     len(self.last_blocks[RiskLevel.HIGH]),
                                                     len(self.last_blocks[RiskLevel.MEDIUM])))

    def get_all_last_risk_block_lst(self):
        if self.last_blocks is None:
            return None
        block_lst = []
        for risk_lv, blocks in self.last_blocks.items():
            block_lst.extend(blocks)
        return block_lst

    def output_risk_record(self):
        print("---------------------统计结果--------------------")
        # 读取上次记录
        self.init_last_risk_block_lst()
        last_risk_blocks = self.get_all_last_risk_block_lst()

        # 写入临时记录文件、标记新增区域、中风险地区按省分类
        f_name = u"{0}RiskBlock_{1}.txt".format(SpiderBase.TEMP_PATH, self.format_result_time_ex())
        f = open(f_name, 'w', encoding="utf8")
        medium_risk_block = {}  # {province : [RiskBlock,]}
        new_blocks = []
        for risk_lv, blocks in self.risk_blocks.items():
            risk_block_count_str = RiskLevel.risk_level_to_text(risk_lv) + ":({0}个)".format(len(blocks))
            f.writelines(risk_block_count_str + "\n")
            # print(risk_block_count_str)
            for block in blocks:
                block_str = block.full_name_with_sep()
                if last_risk_blocks is not None and block not in last_risk_blocks:
                    block.is_new_add = True
                    block_str = block_str + SpiderBase.NEW_ADD_TAG
                    new_blocks.append(block)
                f.writelines(block_str + "\n")
                # print(block_str)

                if risk_lv == RiskLevel.MEDIUM:
                    if block.province not in medium_risk_block:
                        medium_risk_block[block.province] = []
                    medium_risk_block[block.province].append(block)

            f.writelines("\n")
            # print("\n")
        f.close()

        # 调整风险等级的地区
        adjust_blocks = self.calc_adjust_risk_lv_block()
        adjust_high_blocks = adjust_blocks.get(RiskLevel.HIGH)
        adjust_medium_blocks = adjust_blocks.get(RiskLevel.MEDIUM)
        adjust_low_blocks = adjust_blocks.get(RiskLevel.LOW)
        adjust_high_block_str = self.join_block(u"、", adjust_high_blocks)
        adjust_medium_block_str = self.join_block(u"、", adjust_medium_blocks)
        adjust_low_block_str = self.join_block(u"、", adjust_low_blocks)

        # 高风险区列表
        high_risk_str = ""
        high_block_idx = 1
        high_risk_blocks = self.risk_blocks[RiskLevel.HIGH]
        for block in high_risk_blocks:
            high_risk_str += "{0}.{1}\n".format(high_block_idx, block.full_name)
            high_block_idx += 1

        # 中风险区列表
        med_risk_str = ""
        medium_block_idx = 1
        for prov, blocks in medium_risk_block.items():
            med_risk_str += u"{0}（共{1}个）\n".format(prov, len(blocks))
            for block in blocks:
                med_risk_str += "{0}.{1}\n".format(medium_block_idx, block.block_name_with_new_tag())
                medium_block_idx += 1
            # med_risk_str += "\n"

        print("今日({0})高风险区:{1}个, 中风险区:{2}个".format(self.format_result_time(), len(high_risk_blocks),
                                                   medium_block_idx - 1))
        print("新增区域({0}个):\n{1}\n".format(len(new_blocks), self.join_block("\n", new_blocks)))
        print("调为高风险区({0}个):\n{1}\n".format(len(adjust_high_blocks), self.join_block("\n", adjust_high_blocks)))
        print("调为中风险区({0}个):\n{1}\n".format(len(adjust_medium_blocks), self.join_block("\n", adjust_medium_blocks)))
        print("调为低风险区({0}个):\n{1}\n".format(len(adjust_low_blocks), self.join_block("\n", adjust_low_blocks)))
        print("高风险区(共{0}个):".format(len(high_risk_blocks)))
        print(high_risk_str)
        print("中风险区(共{0}个):".format(medium_block_idx - 1))
        print(med_risk_str)

        result_time_str = self.format_result_time()
        # result_time_str = self.result_time.strftime("%m月%d日%H时")
        cur_time_tup = datetime.datetime.now().timetuple()
        cur_date_str = u"{0}年{1}月{2}日".format(cur_time_tup.tm_year, cur_time_tup.tm_mon, cur_time_tup.tm_mday)
        # cur_date_str = time.strftime("%Y年%m月%d日", time.localtime(now_time))

        context = {
            "ResultTime": result_time_str,
            "NewHighRiskBlock": adjust_high_block_str,
            "NewMediumRiskBlock": adjust_medium_block_str,
            "CurDate": cur_date_str,
            "NewLowRiskBlock": adjust_low_block_str,
            "HighRiskBlockCount": str(len(high_risk_blocks)),
            "HighRiskBlockContent": high_risk_str,
            "MediumRiskBlockCount": str(medium_block_idx - 1),
            "MediumRiskBlockContent": med_risk_str,
        }

        self.gen_notice_file(context, medium_risk_block)

    def gen_notice_file(self, context, medium_risk_block):
        # 根据通知模板生成通知
        doc = Document(u'{0}模板.docx'.format(SpiderBase.TEMPLATE_PATH))
        doc.styles.add_style('Song', WD_STYLE_TYPE.CHARACTER).font.name = '宋体'  # 添加字体样式-Song
        doc.styles['Song']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        paragraphs = doc.paragraphs
        for par in paragraphs:
            # 中风险区列表
            for run in par.runs:
                if run.text == "MediumRiskBlockContent":
                    # r = par.add_run("测试Run")
                    preset_font_name = run.font.name
                    preset_font_size = run.font.size
                    doc.styles.add_style('MediumRiskList_Style', WD_STYLE_TYPE.CHARACTER).font.name = preset_font_name  # 添加字体样式
                    doc.styles['MediumRiskList_Style']._element.rPr.rFonts.set(qn('w:eastAsia'), preset_font_name)
                    # print(preset_font)
                    par.clear()
                    medium_block_idx = 1
                    for prov, blocks in medium_risk_block.items():
                        r = par.add_run(u"{0}（共{1}个）\n".format(prov, len(blocks)), style="MediumRiskList_Style")
                        r.font.size = preset_font_size
                        r.bold = True
                        for block in blocks:
                            r = par.add_run("{0}.{1}\n".format(medium_block_idx, block.block_name_with_new_tag()), style="MediumRiskList_Style")
                            r.font.size = preset_font_size
                            r.bold = False
                            medium_block_idx += 1
                    break
                # 其他文本替换
                for k, v in context.items():
                    # print(run.text)
                    run.text = run.text.replace(k, v)
                # if run.text.find("MediumRiskBlockContent") >= 0:
                #     run.text = run.text.replace("MediumRiskBlockContent", "")
                #     medium_block_idx = 1
                #     style = run.style
                #     for prov, blocks in medium_risk_block.items():
                #         med_risk_str = u"{0}（共{1}个）\n".format(prov, len(blocks))
                #         new_run = par.add_run(med_risk_str, style)
                #         new_run.add_break()
                #         for block in blocks:
                #             med_risk_str = "{0}.{1}\n".format(medium_block_idx, block.block_name_with_new_tag())
                #             new_run = par.add_run(med_risk_str, style)
                #             new_run.add_break()
                #             medium_block_idx += 1

        result_time_str = self.format_result_time()
        save_file_name = u"{0}关于发布国内疫情风险等级提示的通知（截至{1}）.docx".format(SpiderBase.OUTPUT_PATH, result_time_str)
        doc.save(save_file_name)
        # self.output_new_add_risk_blocks()

    def join_block(self, sep, blocks):
        block_str = ""
        count = len(blocks)
        for i in range(count):
            block_str += blocks[i].full_name
            if i != count - 1:
                block_str += sep
        return block_str

    # 计算调整风险等级的地区
    def calc_adjust_risk_lv_block(self):
        new_high_block = []
        new_medium_block = []
        new_low_block = []
        adjust_blocks = {
            RiskLevel.HIGH: new_high_block,
            RiskLevel.MEDIUM: new_medium_block,
            RiskLevel.LOW: new_low_block
        }
        if self.last_blocks is None:
            return adjust_blocks
        high_block = self.risk_blocks[RiskLevel.HIGH]
        medium_block = self.risk_blocks[RiskLevel.MEDIUM]
        last_high_block = self.last_blocks.get(RiskLevel.HIGH)
        last_medium_block = self.last_blocks.get(RiskLevel.MEDIUM)
        for block in high_block:
            if block in last_high_block:
                continue
            new_high_block.append(block)

        for block in medium_block:
            if block in last_medium_block:
                continue
            new_medium_block.append(block)

        for block in last_high_block:
            if block in high_block:
                continue
            if block in medium_block:
                new_medium_block.append(block)
                continue
            new_low_block.append(block)

        for block in last_medium_block:
            if block in medium_block:
                continue
            if block in high_block:
                new_high_block.append(block)
                continue
            new_low_block.append(block)

        return adjust_blocks

    def read_record_from_file(self, f_name):
        f = open(f_name, 'r', encoding="utf-8")
        risk_blocks = {}
        for line in f.readlines():
            line = line.replace("\n", "")
            if line == "":
                continue
            line = line.replace(SpiderBase.NEW_ADD_TAG, "")
            if line.find(RiskLevel.MEDIUM_TAG) >= 0:
                cur_risk_level = RiskLevel.MEDIUM
                continue
            if line.find(RiskLevel.HIGH_TAG) >= 0:
                cur_risk_level = RiskLevel.HIGH
                continue
            if cur_risk_level not in risk_blocks:
                risk_blocks[cur_risk_level] = []
            block_lst = line.split(" ")
            risk_block = RiskBlock(block_lst[0], "".join(block_lst[1:]))

            risk_blocks[cur_risk_level].append(risk_block)
        f.close()
        return risk_blocks

    @staticmethod
    def read_record_from_word(f_name):
        risk_blocks = {RiskLevel.HIGH: [], RiskLevel.MEDIUM: []}
        doc = Document(f_name)
        paragraphs = doc.paragraphs
        is_start = False
        h_start_idx = 0
        h_end_idx = 0
        m_start_idx = 0

        for idx in range(len(paragraphs)):
            par = paragraphs[idx]
            idx += 1
            # print(par.text)
            if par.text == "":
                continue
            if par.text.find(u"国内疫情中高风险地区列表") >= 0:
                is_start = True
                continue
            if not is_start:
                continue
            if par.text.find(u"高风险地区") >= 0:
                h_start_idx = idx
                continue
            if par.text.find(u"中风险地区") >= 0:
                m_start_idx = idx
                h_end_idx = idx - 2
                continue

        for h_idx in range(h_start_idx, h_end_idx + 1):
            h_par = paragraphs[h_idx]
            h_text = h_par.text.replace(SpiderBase.NEW_ADD_TAG, "")
            for text in h_text.split("\n"):
                # prov_pos = text.find(u"市")
                # if prov_pos < 0:
                #     prov_pos = text.find(u"区")
                # if prov_pos < 0:
                #     continue
                # province = text[0:prov_pos+1]
                # block = text[prov_pos + 1:]
                # print(province, block)
                # 不分割**省**市，直接用全名
                dot_pos = text.find(u".")
                block_name = text[dot_pos + 1:]
                risk_block = RiskBlock("", block_name)
                risk_blocks[RiskLevel.HIGH].append(risk_block)

        for m_idx in range(m_start_idx, len(paragraphs)):
            m_par = paragraphs[m_idx]
            m_text = m_par.text.replace(SpiderBase.NEW_ADD_TAG, "")
            for text in m_text.split("\n"):
                if not re.match(r"^\d.*", text):
                    # prov_pos = text.find(u"省")
                    # if prov_pos < 0:
                    #     prov_pos = text.find(u"市")
                    prov_pos = text.find(u"（")
                    # if prov_pos < 0:
                    #     continue
                    m_province = text[0:prov_pos]
                else:
                    dot_pos = text.find(u".")
                    if dot_pos >= 0:
                        block_name = text[dot_pos + 1:]
                        risk_block = RiskBlock(m_province, block_name)
                        risk_blocks[RiskLevel.MEDIUM].append(risk_block)
                        # print(m_province, block_name)
        # print("&&&&&&&&&&&&&&&&&&&&&")
        # for block in risk_blocks[RiskLevel.MEDIUM]:
        #     print(block.province, block.block)
        return risk_blocks

    # 输出新增区域
    def output_new_add_risk_blocks(self):
        new_risk_blocks = {}
        str_time = time.strftime("%Y-%m-%d", time.localtime(time.time()))
        f_name = u"{0}RiskBlock_{1}.txt".format(SpiderBase.TEMP_PATH, str_time)
        new_f_name = u"{0}NewRiskBlock_{1}.txt".format(SpiderBase.TEMP_PATH, str_time)

        last_str_time = time.strftime("%Y-%m-%d", time.localtime(time.time() - 3600 * 24))
        last_f_name = u"{0}RiskBlock_{1}.txt".format(SpiderBase.TEMP_PATH, last_str_time)

        if not os.path.exists("last_f_name"):
            print("无法输出新增记录,前一天的记录不存在!!!")
            return

        risk_blocks = self.read_record_from_file(f_name)
        last_risk_blocks = self.read_record_from_file(last_f_name)
        new_f = open(new_f_name, 'w', encoding="utf8")

        for risk_lv, risk_blocks in risk_blocks.items():
            last_blocks = last_risk_blocks.get(risk_lv)
            for risk_block in risk_blocks:
                if last_blocks is None or risk_block not in last_blocks:
                    if risk_lv not in new_risk_blocks:
                        new_risk_blocks[risk_lv] = []
                    new_risk_blocks[risk_lv].append(risk_block)
        for risk_lv, blocks in new_risk_blocks.items():
            new_risk_info = "新增{0}{1}个".format(RiskLevel.risk_level_to_text(risk_lv), len(blocks))
            print(new_risk_info)
            new_f.writelines(new_risk_info)
            for block in blocks:
                print(block)
                new_f.writelines(block)
        new_f.close()

    @staticmethod
    def walk_dir(dir_path, callback=None):
        if not callback:
            return
        list_dir = os.walk(dir_path)
        for (dirPath, _, files) in list_dir:
            for f_name in files:
                callback(dirPath, f_name)


if __name__ == '__main__':
    SpiderBase.read_record_from_word(SpiderBase.OUTPUT_PATH + "关于发布国内疫情风险等级提示的通知（截至8月14日13时）.docx")
